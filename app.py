import streamlit as st
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from langchain_community.document_loaders import WebBaseLoader
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_huggingface import HuggingFaceEndpoint
from secret_api_keys import huggingface_api_key
from streamlit.runtime.uploaded_file_manager import UploadedFile

def process_input(input_type, input_data):
    '''Process different types of input data and returns a vector store.'''
    loader = None
    if input_type == "Link":
        loader = WebBaseLoader(input_data)
        documents = loader.load()
    elif input_type == "PDF":
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input type for PDF.")
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        documents = text

    elif input_type == "DOCX":
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input type for DOCX.")
        text = "\n".join([para.text for para in doc.paragraphs])
        documents = text

    elif input_type == "Text":
        if isinstance(input_data, str):
            documents = input_data
        else:
            raise ValueError("Invalid input type for Text.")
        
    elif input_type == "TXT":
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = str(input_data.read().decode('utf-8'))
        else:
            raise ValueError("Invalid input type for TXT.")
        documents = text

    else:
        raise ValueError("Unsupported input type.")
    
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    if input_type == "Link":
        texts = text_splitter.split_documents(documents)
        texts = [str(doc.page_content) for doc in texts]
    else:
        texts = text_splitter.split_text(documents)

    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )

    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)

    vector_store = FAISS(
        index=index,
        embedding_function=hf_embeddings.embed_query,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={}
    )
    vector_store.add_texts(texts)
    return vector_store


def answer_question(vectorstore, query):
    '''Answer a question using the vector store.'''
    if not hasattr(vectorstore, "index_to_docstore_id") or not vectorstore.index_to_docstore_id:
        return "No documents found in the vector store. Please upload or enter some data first."
    retriever = vectorstore.as_retriever()
    docs = retriever.get_relevant_documents(query)
    if not docs:
        return "No relevant documents found for your query."
    llm = HuggingFaceEndpoint(
        repo_id='microsoft/DialoGPT-medium',
        huggingfacehub_api_token=huggingface_api_key,
        temperature=0.6
    )
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=vectorstore.as_retriever())
    answer = qa.invoke({"query": query})
    return answer


def simple_answer_question(vectorstore, query):
    '''Fixed version that handles parameter validation correctly'''
    try:
        print(f"DEBUG: Starting with query: {query}")
        
        if not vectorstore or not hasattr(vectorstore, "index_to_docstore_id"):
            return "Vector store not available"
        
        # Get documents
        retriever = vectorstore.as_retriever(search_kwargs={"k": 2})
        docs = retriever.get_relevant_documents(query)
        print(f"DEBUG: Retrieved {len(docs)} documents")
        
        if not docs:
            return "No relevant documents found"
        
        # Create context
        context = docs[0].page_content[:400] if docs else ""
        print(f"DEBUG: Context length: {len(context)} characters")
        
        # Fixed LLM configuration
        try:
            print("DEBUG: Creating LLM with correct parameters...")
            
            llm = HuggingFaceEndpoint(
                repo_id='microsoft/DialoGPT-medium',
                huggingfacehub_api_token=huggingface_api_key,
                temperature=0.6,
                max_new_tokens=50,  # Move this out of model_kwargs
                model_kwargs={
                    "pad_token_id": 50256,
                    "do_sample": True
                    # Remove max_new_tokens and max_length from here
                }
            )
            print("DEBUG: LLM created successfully")
            
            # Create a simple prompt
            prompt = f"{context[:200]}\n\nQuestion: {query}\nAnswer:"
            print(f"DEBUG: Prompt created, length: {len(prompt)}")
            
            print("DEBUG: Calling LLM...")
            response = llm.invoke(prompt)
            print(f"DEBUG: Response received: {type(response)}")
            
            if response:
                response_str = str(response).strip()
                print(f"DEBUG: Response string: '{response_str[:100]}...'")
                
                if response_str and len(response_str) > 5:
                    return f"AI Response: {response_str}"
                else:
                    return f"Empty AI response. Here's the relevant info: {context[:300]}"
            else:
                return f"No response from AI. Here's the relevant info: {context[:300]}"
                
        except Exception as llm_error:
            print(f"DEBUG: LLM Error Type: {type(llm_error).__name__}")
            print(f"DEBUG: LLM Error Message: {str(llm_error)}")
            
            # Try with even simpler parameters
            try:
                print("DEBUG: Trying with minimal parameters...")
                simple_llm = HuggingFaceEndpoint(
                    repo_id='microsoft/DialoGPT-medium',
                    huggingfacehub_api_token=huggingface_api_key,
                    max_new_tokens=30  # Only essential parameters
                )
                
                short_prompt = f"Context: {context[:150]}\nQ: {query}\nA:"
                response = simple_llm.invoke(short_prompt)
                
                if response:
                    return f"Simple AI Response: {str(response).strip()}"
                else:
                    return f"Here's what I found about '{query}': {context[:400]}"
                    
            except Exception as simple_error:
                print(f"DEBUG: Simple LLM also failed: {str(simple_error)}")
                return f"AI unavailable. Here's relevant information about '{query}':\n\n{context[:800]}"
            
    except Exception as outer_error:
        print(f"DEBUG: Outer error: {str(outer_error)}")
        return f"Error: {str(outer_error)}"


def minimal_answer_question(vectorstore, query):
    '''Most minimal version to avoid all parameter issues'''
    try:
        # Get documents
        retriever = vectorstore.as_retriever(search_kwargs={"k": 2})
        docs = retriever.get_relevant_documents(query)
        
        if not docs:
            return "No relevant documents found"
        
        context = docs[0].page_content[:400]
        
        try:
            # Try with absolute minimal parameters
            llm = HuggingFaceEndpoint(
                repo_id='gpt2',  # Use GPT-2 which is very reliable
                huggingfacehub_api_token=huggingface_api_key
                # No other parameters at all
            )
            
            prompt = f"Answer based on this: {context[:200]}\nQuestion: {query}\nAnswer:"
            response = llm.invoke(prompt)
            
            if response and str(response).strip():
                return str(response).strip()
            else:
                return f"Based on the documents about '{query}': {context}"
                
        except:
            # If even this fails, just return the context
            return f"Here's information about '{query}' from the documents:\n\n{context}"
            
    except Exception as e:
        return f"Error: {str(e)}"


def main():
    st.title("Advanced RAG-QnA Application")
    input_type = st.selectbox("Select Input Type", ["Link", "PDF", "DOCX", "Text", "TXT"])
    if input_type == "Link":
        number_input = st.number_input("Enter Number of Links", min_value=1, max_value=20, step=1)
        input_data = []
        for i in range(number_input):
            url = st.sidebar.text_input(f"Link {i + 1}")
            if url:
                input_data.append(url)
    elif input_type == "PDF":
        input_data = st.file_uploader("Upload PDF files", type=["pdf"])
    elif input_type == "DOCX":  
        input_data = st.file_uploader("Upload DOCX files", type=["docx", "doc"])
    elif input_type == "Text":
        input_data = st.text_area("Enter Text")
    elif input_type == "TXT":
        input_data = st.file_uploader("Upload TXT files", type=["txt"])
    
    # if st.button("Proceed"):
    #     vectorstore = process_input(input_type, input_data)
    #     st.session_state["vectorstore"] = vectorstore

    if st.button("Proceed"):
        if input_type == "Link":
            if not input_data or all(not url.strip() for url in input_data):
                st.error("Please enter at least one valid link.")
            else:
                vectorstore = process_input(input_type, input_data)
                st.session_state["vectorstore"] = vectorstore
        elif input_type in ["PDF", "DOCX", "TXT"]:
            if input_data is None:
                st.error(f"Please upload a {input_type} file.")
            else:
                vectorstore = process_input(input_type, input_data)
                st.session_state["vectorstore"] = vectorstore
        elif input_type == "Text":
            if not input_data or not input_data.strip():
                st.error("Please enter some text.")
            else:
                vectorstore = process_input(input_type, input_data)
                st.session_state["vectorstore"] = vectorstore
    
    if "vectorstore" in st.session_state:
        query = st.text_input("Ask a question")
        if st.button("Submit"):
            answer = simple_answer_question(st.session_state["vectorstore"], query)
            st.write(answer)


if __name__ == "__main__":
    main()